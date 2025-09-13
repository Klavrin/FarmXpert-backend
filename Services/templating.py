# Services/templating.py
from __future__ import annotations
import re, io, zipfile, pathlib
from typing import Dict, Any, List
from docx import Document as DocxDocument
from openpyxl import load_workbook

LABEL_MAP = {
  r"Denumirea solicitantului": "users.fullName",
  r"Telefon": "users.phone",
  r"E[-]?mail": "users.email",
}

def _value_from_dataset(dataset: Dict[str,Any], key: str) -> str | None:
    # key e.g., users.fullName
    table, attr = key.split(".",1)
    if table == "users" and dataset.get("users"):
        u = dataset["users"][0]
        if attr == "fullName":
            return f"{u.get('firstName','')} {u.get('lastName','')}".strip() or None
        return u.get(attr)
    return None

def fill_docx(input_path: pathlib.Path, dataset: Dict[str,Any], out_path: pathlib.Path):
    with open(input_path, "rb") as f:
        data = f.read()
    doc = DocxDocument(io.BytesIO(data))

    # 1) {{placeholders}}
    placeholder_re = re.compile(r"\{\{([^}]+)\}\}")
    for p in doc.paragraphs:
        m = placeholder_re.search(p.text)
        if m:
            key = m.group(1).strip()
            val = _value_from_dataset(dataset, key) or ""
            p.text = placeholder_re.sub(val, p.text)

    # 2) Replace underscores after labels
    for p in doc.paragraphs:
        txt = p.text
        for label, key in LABEL_MAP.items():
            if re.search(label, txt, re.I) and re.search(r"_\s*[_]+", txt):
                val = _value_from_dataset(dataset, key) or ""
                txt = re.sub(r"(_\s*[_]+)", f" {val} ", txt)
        p.text = txt

    doc.save(out_path)

def fill_xlsx(input_path: pathlib.Path, dataset: Dict[str,Any], out_path: pathlib.Path):
    wb = load_workbook(filename=str(input_path))
    for ws in wb.worksheets:
        # simple header→value fill on left column labels
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 100), max_col=min(ws.max_column, 10)):
            for cell in row:
                if isinstance(cell.value, str):
                    # try known labels
                    for label, key in LABEL_MAP.items():
                        if re.fullmatch(label, cell.value.strip(), flags=re.I):
                            v = _value_from_dataset(dataset, key) or ""
                            # write in next cell (to the right) if empty
                            tgt = ws.cell(row=cell.row, column=cell.column+1)
                            if tgt.value in (None, "", "Unnamed: 1"):
                                tgt.value = v
    wb.save(str(out_path))

def package_zip(files: List[pathlib.Path], zip_path: pathlib.Path):
    with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as z:
        for p in files:
            z.write(str(p), arcname=p.name)
