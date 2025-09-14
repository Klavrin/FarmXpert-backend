# app/scraper/routes.py
import os, json, pathlib, re, io
import uuid, base64, pathlib, requests
from flask import Blueprint, request, jsonify, send_file
from sqlalchemy import text
from app.db_utilis import engine
from docx import Document

# helper to copy basic run formatting
def _copy_run_formatting(src_run, dst_run):
    try:
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
    except Exception:
        pass
    try:
        if src_run.font.name:
            dst_run.font.name = src_run.font.name
    except Exception:
        pass
    try:
        if src_run.font.size:
            dst_run.font.size = src_run.font.size
    except Exception:
        pass
    try:
        if src_run.font.color and src_run.font.color.rgb:
            dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass


def _norm_doc_type(s: str | None) -> str | None:
    if not s:
        return None
    t = (s or "").strip().lower()
    # unify common variants (keep it simple)
    if t.startswith("cerere"): return "cerere"
    if "fisa" in t or "fișa" in t or "fişă" in t: return "fișă de calcul"
    if t.startswith("act"): return "act"
    if t.startswith("anex"): return "anexă"
    if t.startswith("ordin"): return "ordin"
    if "hotar" in t or "hotăr" in t: return "hotărâre"
    if t.startswith("ghid"): return "ghid"
    if "angajament" in t: return "contract"
    return t

def _fetch_docx_bytes(url: str) -> bytes:
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    # be strict about extension; many servers mislabel content-type
    if not url.lower().endswith(".docx"):
        raise ValueError("Resolved URL is not a .docx")
    return r.content

def _docx_to_text(doc_bytes: bytes) -> str:
    doc = Document(io.BytesIO(doc_bytes))
    chunks = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    chunks.append(cell.text)
    text = "\n".join(chunks)
    # squeeze long whitespace
    return "\n".join([ln.strip() for ln in text.splitlines() if ln.strip()])

def _call_ai_transform(input_text: str, instructions: str | None, lang: str = "ro") -> str:
    """
    Calls an LLM to transform/clean the DOCX text.
    Works with OpenAI v1.x (preferred) and falls back to old sdk if present.
    Env: OPENAI_API_KEY; optional OPENAI_MODEL (default: gpt-4o-mini or gpt-3.5-turbo)
    """
    sys = (
        "You rewrite/transform the provided document text for a .docx output. "
        "Keep it factual, avoid hallucinations, preserve lists as lists, and output PLAIN TEXT only.\n"
        "If there are form-like fields, keep clear headings and bullet points."
    )
    user = (
        f"Language: {lang}\n"
        f"Instructions: {instructions or 'Clean up and structure for a downloadable .docx.'}\n\n"
        f"--- SOURCE TEXT START ---\n{input_text}\n--- SOURCE TEXT END ---\n"
        "Return ONLY the rewritten text. No markdown, no code fences."
    )
    model_new = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    model_old = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")

    # Try new SDK first
    try:
        from openai import OpenAI
        client = OpenAI()
        resp = client.chat.completions.create(
            model=model_new,
            messages=[{"role":"system","content":sys},{"role":"user","content":user}],
            temperature=0.2,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception:
        pass

    # Fallback to old SDK
    try:
        import openai
        openai.api_key = os.getenv("OPENAI_API_KEY")
        resp = openai.ChatCompletion.create(
            model=model_old,
            messages=[{"role":"system","content":sys},{"role":"user","content":user}],
            temperature=0.2,
        )
        return (resp["choices"][0]["message"]["content"] or "").strip()
    except Exception as e:
        raise RuntimeError(f"AI call failed: {e}")

def _write_docx_from_text(text: str) -> bytes:
    doc = Document()
    # simple rule: blank line → paragraph break, bullets preserved by leading "- " or "• "
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")  # blank line
            continue
        if line.lstrip().startswith(("- ", "• ")):
            p = doc.add_paragraph()
            p.style = doc.styles["List Bullet"] if "List Bullet" in doc.styles else None
            p.add_run(line.lstrip()[2:].strip())
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# 0-------------------------------------------------------------------------
# Helpers for autocomplete
_BLANK_RE = re.compile(r"(?:_{4,}|\.{4,}|_{2,}\.{2,}|\.{2,}_{2,})")

def _slug(s: str) -> str:
    s = re.sub(r"\s+", "_", (s or "").strip().lower())
    s = re.sub(r"[^a-z0-9_]", "", s)
    return s or "field"

def _guess_label(text: str, match: re.Match) -> str:
    # Look left of the blank for a label or parenthetical hint
    left = text[:match.start()]
    # parenthetical (...) just before the blank?
    m = re.search(r"\(([^\)]+)\)\s*$", left)
    if m:
        return m.group(1).strip()
    # label: ______  OR  label ______
    m = re.search(r"([A-Za-zăâîșțA-Z0-9/ ,\-]{3,})[: ]\s*$", left)
    if m:
        return m.group(1).strip()
    # fallback: last few words
    tokens = re.findall(r"[A-Za-zăâîșț0-9]+", left)[-4:]
    return " ".join(tokens).strip() or "câmp"

def _extract_fields_from_paragraph_text(text: str):
    fields = []
    for m in _BLANK_RE.finditer(text):
        label = _guess_label(text, m)
        context = text[max(0, m.start()-80): m.end()+40]
        fields.append({"id": _slug(label), "label": label, "context": context})
    return fields

def _extract_doc_fields(doc: Document):
    fields = []
    # paragraphs
    for p in doc.paragraphs:
        fields.extend(_extract_fields_from_paragraph_text(p.text))
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                fields.extend(_extract_fields_from_paragraph_text(cell.text))
    # de-dup by id while keeping order
    seen = set()
    uniq = []
    for f in fields:
        k = f["id"]
        if k in seen: 
            continue
        seen.add(k)
        uniq.append(f)
    return uniq

def _ai_fill_fields(fields, profile=None, instructions=None, language="ro"):
    """
    Returns a dict {field_id: value} using an AI provider if configured,
    else generates reasonable placeholders.
    """
    # 1) OpenAI (optional) if OPENAI_API_KEY is present
    try:
        import os
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)

            # Keep it robust & cheap; ask for strict JSON back.
            sys = (
                "Ești un asistent care completează formulare agricole din R. Moldova. "
                "Întoarce STRICT un JSON object cu chei EXACT câmpurile primite."
            )
            user_msg = {
                "language": language,
                "instructions": instructions or "Completează realist, concis și formal.",
                "profile": profile or {},
                "fields": [{"id": f["id"], "label": f["label"], "context": f["context"]} for f in fields],
            }
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": sys},
                    {"role": "user", "content": json.dumps(user_msg, ensure_ascii=False)},
                    {"role": "user", "content": "Returnează DOAR JSON-ul cu {id: valoare}."},
                ],
                temperature=0.2,
            )
            raw = resp.choices[0].message.content
            data = json.loads(raw)
            # ensure only requested keys
            out = {}
            want = {f["id"] for f in fields}
            for k, v in (data or {}).items():
                if k in want:
                    out[k] = str(v)
            # fill any missing keys with placeholders
            for f in fields:
                out.setdefault(f["id"], f"[completați: {f['label']}]")
            return out
    except Exception:
        # fall through to local
        pass

    # 2) Local fallback (no external API)
    out = {}
    profile = profile or {}
    for f in fields:
        label = f["label"]
        # simple heuristic matches profile keys
        key_guess = _slug(label)
        v = profile.get(key_guess) or profile.get(label) or f"[exemplu automat pentru: {label}]"
        out[f["id"]] = str(v)
    return out

def _jaccard(a: str, b: str) -> float:
    ta = set(re.findall(r"[a-z0-9ăâîșț]+", (a or "").lower()))
    tb = set(re.findall(r"[a-z0-9ăâîșț]+", (b or "").lower()))
    return 0.0 if not ta or not tb else len(ta & tb) / len(ta | tb)

def _best_suggestion_for(label: str, suggestions: dict):
    # pick best by Jaccard label similarity on suggestion keys
    best_k, best_s = None, 0.0
    for k in suggestions.keys():
        s = _jaccard(label, k)
        if s > best_s:
            best_k, best_s = k, s
    return best_k

def _apply_suggestions_inline(doc: Document, suggestions: dict):
    """
    For each blank (____ or .....), append a suggestion right after it as «value».
    Safer rewriting: we preserve the first run formatting for paragraphs/cells where we change text.
    """
    def replace_text(txt: str):
        out, idx = [], 0
        for m in _BLANK_RE.finditer(txt):
            out.append(txt[idx:m.end()])  # keep the blank itself
            label = _guess_label(txt, m)
            key = _best_suggestion_for(_slug(label), suggestions) or _best_suggestion_for(label, suggestions)
            val = suggestions.get(key) if key else None
            if val:
                out.append(f" «{val}»")
            idx = m.end()
        out.append(txt[idx:])
        return "".join(out)

    # paragraphs
    for p in list(doc.paragraphs):
        original = p.text or ""
        new_text = replace_text(original)
        if new_text != original:
            # preserve first run formatting where possible
            first_run = p.runs[0] if p.runs else None
            # clear all run text (safe approach)
            for run in list(p.runs):
                try:
                    run.text = ""
                except Exception:
                    pass
            # add a single run with new_text
            new_run = p.add_run(new_text)
            if first_run:
                try:
                    _copy_run_formatting(first_run, new_run)
                except Exception:
                    pass

    # tables (cells) — treat each cell's first paragraph similarly
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if not cell.paragraphs:
                    continue
                # join original cell text and produce replacement; this is conservative
                cell_text = cell.text or ""
                new_text = replace_text(cell_text)
                if new_text != cell_text:
                    # capture first paragraph's first run formatting (if any)
                    first_para = cell.paragraphs[0]
                    first_run = first_para.runs[0] if first_para.runs else None
                    # clear all paragraphs in the cell and recreate single paragraph
                    for para in list(cell.paragraphs):
                        for run in list(para.runs):
                            try:
                                run.text = ""
                            except Exception:
                                pass
                        # remove extra paragraphs by setting their text empty (python-docx doesn't provide direct remove)
                        try:
                            para.text = ""
                        except Exception:
                            pass
                    # write into first paragraph
                    target_para = cell.paragraphs[0]
                    try:
                        # clear existing runs on target para
                        for rr in list(target_para.runs):
                            rr.text = ""
                    except Exception:
                        pass
                    new_run = target_para.add_run(new_text)
                    if first_run:
                        try:
                            _copy_run_formatting(first_run, new_run)
                        except Exception:
                            pass


def _resolve_source_to_url(data, conn):
    """
    Resolve source .docx URL from:
      - direct "url"
      - or {code, doc_type} in scraped_asset
    """
    url = (data.get("url") or "").strip()
    if url:
        return url

    code = (data.get("code") or "").strip().upper()
    doc_type = (data.get("doc_type") or "").strip().lower() or None

    if not code:
        return None

    where = ["measure_code = :code", "ext in ('docx','DOCX')"]
    params = {"code": code}
    if doc_type:
        where.append("lower(coalesce(doc_type,'')) like :dt")
        params["dt"] = doc_type + "%"

    sql = f"""
      select url
      from public.scraped_asset
      where {' AND '.join(where)}
      order by
        case when lower(coalesce(doc_type,'')) like 'cerere%%' then 0
             when lower(coalesce(doc_type,'')) like 'fișă%%' or lower(coalesce(doc_type,'')) like 'fisa%%' then 1
             when lower(coalesce(doc_type,'')) like 'act%%' then 2
             else 9
        end,
        length(coalesce(filename,'')) asc
      limit 1
    """
    row = conn.execute(text(sql), params).first()
    return row[0] if row else None

def _download_bytes(url: str) -> bytes:
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    return r.content

def _bootstrap_autocomplete(conn):
    # Try to enable helpful extensions (safe to re-run)
    for ext in ("unaccent", "pg_trgm"):
        try:
            conn.execute(text(f"create extension if not exists {ext}"))
        except Exception:
            # no perms? ignore; autocomplete will still work (just slower/fuzzier)
            pass

    # 1) Ensure a regular tsvector column exists (we avoid generated column with unaccent,
    #    because generated expressions require IMMUTABLE functions)
    try:
        conn.execute(text("""
            alter table public.scraped_asset
            add column if not exists search tsvector
        """))
    except Exception:
        # ignore if something odd happens (race/perm), caller will still attempt to continue
        pass

    # 2) Create or replace trigger function that populates `search`.
    #    Use unaccent + to_tsvector inside the trigger (allowed).
    try:
        conn.execute(text("""
            create or replace function public.scraped_asset_search_update() returns trigger as $$
            begin
              -- Build accent-insensitive tsvector from filename, doc_type and summary
              new.search := to_tsvector(
                'simple',
                coalesce(unaccent(new.filename), '') || ' ' ||
                coalesce(unaccent(new.doc_type), '') || ' ' ||
                coalesce(unaccent(new.summary), '')
              );
              return new;
            end;
            $$ language plpgsql;
        """))
    except Exception:
        # If creating the function fails (permissions), continue; searches may be less performant
        pass

    # 3) Drop & create the trigger so it's idempotent
    try:
        conn.execute(text("drop trigger if exists trg_scraped_asset_search on public.scraped_asset"))
    except Exception:
        pass

    try:
        conn.execute(text("""
            create trigger trg_scraped_asset_search
            before insert or update on public.scraped_asset
            for each row execute procedure public.scraped_asset_search_update();
        """))
    except Exception:
        # if trigger creation fails (e.g. no permissions), keep going
        pass

    # 4) Indexes (idempotent). The trigram index uses unaccent in expression; wrap in try/except.
    try:
        conn.execute(text("""
            create index if not exists idx_scraped_asset_search
            on public.scraped_asset using gin (search)
        """))
    except Exception:
        pass

    try:
        conn.execute(text("""
            create index if not exists idx_scraped_asset_trgm
            on public.scraped_asset using gin (unaccent(filename) gin_trgm_ops)
        """))
    except Exception:
        # some DBs / users won't have permission to create expression indexes; ignore
        pass

    try:
        conn.execute(text("""
            create index if not exists idx_scraped_asset_measure
            on public.scraped_asset (measure_code)
        """))
    except Exception:
        pass

    # 5) Optional convenience view (safe to re-run)
    try:
        conn.execute(text("""
            create or replace view public.subsidy_doc_autocomplete as
            select
              a.url, a.filename, a.doc_type, a.ext, a.language,
              a.measure_code as code,
              a.summary,
              s.id as subsidy_id,
              s.title as subsidy_title
            from public.scraped_asset a
            left join public.subsidy s on s.code = a.measure_code
        """))
    except Exception:
        # ignore view creation errors (permissions, etc.)
        pass


MEASURE_RE = re.compile(r'\b(SP|SC)[ _.\-]?(\d+(?:\.\d+)*)\b', re.IGNORECASE)

def _infer_doc_type(filename: str | None, fallback: str | None = None) -> str | None:
    name = (filename or "").lower()
    if "cerere" in name: return "cerere"
    if "fișa de calcul" in name or "fisa de calcul" in name or "fișa" in name: return "fișă de calcul"
    if name.startswith("act ") or "act " in name: return "act"
    if "anexa" in name or "anexă" in name: return "anexă"
    if "angajament" in name: return "contract"
    if "ordin" in name: return "ordin"
    if "hotărâre" in name or "hotarare" in name or name.startswith("hg"): return "hotărâre"
    if "ghid" in name: return "ghid"
    return fallback

def _infer_measure(s: str | None) -> str | None:
    if not s:
        return None
    m = MEASURE_RE.search(s)
    if not m:
        return None
    # Normalize to e.g. "SP_2.10" or "SC_5.4"
    return f"{m.group(1).upper()}_{m.group(2)}"

def _to_text_summary(summary_field):
    """Return (summary_text, language, doc_type, summary_json_or_None)."""
    if isinstance(summary_field, dict):
        return (
            summary_field.get("about") or None,
            summary_field.get("language"),
            summary_field.get("doc_type"),
            summary_field,
        )
    if isinstance(summary_field, str):
        return summary_field or None, None, None, None
    return None, None, None, None

CODE_RE = re.compile(r'\b(SP|SC)[ _.\-]?(\d+(?:\.\d+)*)\b', re.IGNORECASE)

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").lower()).strip()

def _extract_code(*strings):
    for s in strings:
        if not s:
            continue
        m = CODE_RE.search(str(s))
        if m:
            return f"{m.group(1).upper()}_{m.group(2)}"
    return None

def _jaccard(a: str, b: str) -> float:
    ta = set(re.findall(r"[a-z0-9ăâîșț]+", _norm(a)))
    tb = set(re.findall(r"[a-z0-9ăâîșț]+", _norm(b)))
    if not ta or not tb: 
        return 0.0
    return len(ta & tb) / len(ta | tb)

def _persist_summaries(engine, results: list[dict], lang="ro"):
    if not results:
        return {"updated": 0, "inserted": 0, "unmatched": []}

    with engine.begin() as conn:
        # 1) ensure additive column
        conn.execute(text('alter table public.subsidy add column if not exists summary text'))

        # 2) load existing subsidies
        subs = conn.execute(text("""
            select code, title, coalesce(summary,'') as summary
            from public.subsidy
        """)).mappings().all()
        by_code = {r["code"]: dict(r) for r in subs}
        all_titles = [(r["code"], r["title"] or "") for r in subs]

        updated = 0
        inserted = 0
        unmatched = []

        for item in results:
            url   = item.get("url", "")
            fn    = item.get("filename", "")
            prev  = item.get("text_preview", "") or ""
            summ  = (item.get("summary") or "").strip()
            if not summ:
                continue

            # 3) try direct code
            code = _extract_code(fn, url, prev)
            title_guess = re.sub(r"[_-]+", " ", re.sub(r"\.[a-z0-9]+$", "", fn, flags=re.I)).strip()

            matched_code = None
            # Known code?
            if code and code in by_code:
                matched_code = code
            # Unknown code -> seed minimal row
            elif code:
                conn.execute(text("""
                    insert into public.subsidy (code, title, language, status, created_at)
                    values (:c, :t, :lang, 'activ', now())
                    on conflict (code) do nothing
                """), {"c": code, "t": title_guess or code, "lang": lang})
                by_code.setdefault(code, {"code": code, "title": title_guess, "summary": ""})
                matched_code = code
                inserted += 1
            else:
                # 4) fuzzy match by title/filename
                hay = f"{title_guess} {url}"
                best = None
                best_s = 0.0
                for c, t in all_titles:
                    s = max(_jaccard(hay, t), _jaccard(prev[:400], t))  # look at filename/url and a slice of text
                    if s > best_s:
                        best_s, best = s, c
                if best and best_s >= 0.35:
                    matched_code = best

            if not matched_code:
                unmatched.append({"filename": fn, "url": url})
                continue

            # 5) update summary only if an improvement
            row = by_code.get(matched_code, {"summary": ""})
            old = (row.get("summary") or "").strip()
            if len(summ) > len(old):
                conn.execute(text("""
                    update public.subsidy
                    set summary = :s
                    where code = :c
                """), {"s": summ, "c": matched_code})
                by_code[matched_code]["summary"] = summ
                updated += 1

    return {"updated": updated, "inserted": inserted, "unmatched": unmatched}

# Keep your original import path casing
from Services.scrape import scrape_and_summarize, discover_file_links, DEFAULT_EXTS

scraper_bp = Blueprint("scraper", __name__, url_prefix="/api/scraper")
bp = Blueprint("main", __name__)

URL_SCHEME_RE = re.compile(r"^https?://", re.IGNORECASE)

def _normalize_pages(pages_raw):
    """
    Accepts a single string or a list of strings. Returns (pages, errors)
    where pages is a cleaned list of http(s) URLs.
    """
    errors = []
    if isinstance(pages_raw, str):
        pages = [pages_raw]
    elif isinstance(pages_raw, (list, tuple)):
        pages = [p for p in pages_raw if isinstance(p, str)]
    else:
        return [], ["'pages' must be a string or an array of strings"]

    # trim empties
    pages = [p.strip() for p in pages if p and p.strip()]

    # auto-prepend https:// for bare domains (optional; comment out if you prefer strict 400)
    fixed = []
    for p in pages:
        if URL_SCHEME_RE.match(p):
            fixed.append(p)
        elif re.match(r"^[a-z0-9.-]+\.[a-z]{2,}(/.*)?$", p, re.I):
            fixed.append("https://" + p)
        else:
            errors.append(f"Invalid URL (must start with http:// or https://): {p}")
    return fixed, errors

@bp.route("/")
def home():
    return "Hello, World!"

@scraper_bp.post("/links")
def list_links():
    data = request.get_json(force=True) or {}
    pages_raw = data.get("pages")
    # accept both "ext" and "exts"
    exts = data.get("ext") or data.get("exts") or DEFAULT_EXTS

    pages, errs = _normalize_pages(pages_raw)
    if errs:
        return jsonify(error="Bad 'pages' input", details=errs), 400
    if not pages:
        return jsonify(error="Provide pages: []"), 400

    links, errors = [], []
    for p in pages:
        try:
            links.extend(discover_file_links(p, exts))
        except Exception as e:
            errors.append({"page": p, "error": str(e)})

    links = list(dict.fromkeys(links))
    return jsonify({"pages": pages, "ext": exts, "links": links, "errors": errors})

@scraper_bp.post("/run")
def run_pipeline():
    data = request.get_json(force=True) or {}
    pages_raw = data.get("pages")
    exts = data.get("ext") or data.get("exts") or DEFAULT_EXTS
    lang = data.get("lang", "ro")
    save = data.get("save_dir")
    dry  = bool(data.get("dry_run", False))
    persist = bool(data.get("persist", True))  # default ON

    pages, errs = _normalize_pages(pages_raw)
    if errs:
        return jsonify(error="Bad 'pages' input", details=errs), 400
    if not pages:
        return jsonify(error="Provide pages: []"), 400

    out = scrape_and_summarize(pages, exts=exts, lang=lang, save_dir=save, dry_run=dry)

    if persist and out.get("results"):
        try:
            with engine.begin() as conn:
                # Ensure the raw inbox table exists (safe to re-run)
                conn.execute(text("""
                    create table if not exists public.scraped_asset (
                      id            bigserial primary key,
                      url           text not null unique,
                      filename      text,
                      ext           text,
                      doc_type      text,
                      language      text,
                      summary       text,
                      summary_json  jsonb,
                      text_preview  text,
                      measure_code  text,
                      created_at    timestamptz default now(),
                      updated_at    timestamptz default now()
                    )
                """))
                conn.execute(text("create index if not exists idx_scraped_asset_measure on public.scraped_asset (measure_code)"))
                conn.execute(text("create index if not exists idx_scraped_asset_doc_type on public.scraped_asset (doc_type)"))
                # Optional: make sure subsidy has a summary column
                try:
                    conn.execute(text("alter table public.subsidy add column if not exists summary text"))
                except Exception:
                    pass

                upsert_sql = text("""
                    insert into public.scraped_asset
                      (url, filename, ext, doc_type, language, summary, summary_json, text_preview, measure_code, updated_at)
                    values
                      (:url, :filename, :ext, :doc_type, :language, :summary, :summary_json, :text_preview, :measure_code, now())
                    on conflict (url) do update set
                      filename = excluded.filename,
                      ext = excluded.ext,
                      doc_type = excluded.doc_type,
                      language = excluded.language,
                      summary = excluded.summary,
                      summary_json = excluded.summary_json,
                      text_preview = excluded.text_preview,
                      measure_code = excluded.measure_code,
                      updated_at = now()
                    returning (xmax = 0) as inserted
                """)

                assets_inserted = assets_updated = assets_skipped = 0
                subsidy_updated = 0

                for item in out["results"]:
                    try:
                        url = item.get("url") or item.get("link")
                        if not url:
                            assets_skipped += 1
                            continue

                        filename = item.get("filename")
                        ext = item.get("ext")
                        text_preview = item.get("text_preview") or ""

                        # Gracefully handle dict OR string summaries
                        summary_text, lang_from_summary, doc_type_from_summary, summary_json = _to_text_summary(item.get("summary"))

                        doc_type = doc_type_from_summary or _infer_doc_type(filename, None)
                        language = lang_from_summary
                        if not language:
                            # simple heuristic: Cyrillic => ru, else ro
                            sample = f"{text_preview} {filename or ''}"
                            language = "ru" if re.search(r"[А-Яа-яЁё]", sample) else "ro"

                        measure = (
                            _infer_measure(filename) or
                            _infer_measure(text_preview) or
                            _infer_measure(url) or
                            None
                        )

                        res = conn.execute(upsert_sql, {
                            "url": url,
                            "filename": filename,
                            "ext": ext,
                            "doc_type": doc_type,
                            "language": language,
                            "summary": summary_text,
                            "summary_json": json.dumps(summary_json, ensure_ascii=False) if summary_json else None,
                            "text_preview": text_preview,
                            "measure_code": measure,
                        })
                        inserted = list(res)[0][0]
                        if inserted:
                            assets_inserted += 1
                        else:
                            assets_updated += 1

                        # Keep your old behavior: update subsidy.summary when we can infer a code
                        # NOTE: use the safe summary_text (no .strip() on dicts)
                        code = _extract_code(filename, url, text_preview)  # assuming you already have this helper
                        if code and summary_text:
                            conn.execute(text("""
                                update public.subsidy
                                   set summary = :s
                                 where code = :c
                                   and (summary is null or length(summary) < length(:s))
                            """), {"s": summary_text, "c": code})
                            subsidy_updated += 1

                    except Exception:
                        assets_skipped += 1  # ignore bad rows, continue

                out["persist"] = {
                    "assets_inserted": assets_inserted,
                    "assets_updated": assets_updated,
                    "assets_skipped": assets_skipped,
                    "subsidy_updated": subsidy_updated,
                }

        except Exception as e:
            out["persist"] = {"error": str(e)}

    return jsonify(out), 200

_DOC_ORDER = ["cerere", "fișă de calcul", "fisa de calcul", "act", "anexă", "ordin", "hotărâre", "ghid", "contract", "altele"]

def _normalize_measure(m: str | None) -> str | None:
    if not m:
        return None
    m = m.strip().upper()
    # keep only letters, digits and dots
    m = re.sub(r"[^0-9A-Z. ]", "", m)
    # turn "SP2.1" / "SC 5.4" into "SP_2.1" / "SC_5.4"
    m = re.sub(r'^(SP|SC)\s*', r'\1_', m)
    # ensure the prefix is valid
    if not re.match(r'^(SP|SC)_\d', m):
        return None
    return m

@scraper_bp.get("/autocomplete")
def autocomplete_docs():
    """
    Query params:
      q      - free text (filename/doc type/summary)
      code   - measure code (SP_2.1, SP 2.1, sp2.1 all ok)
      lang   - 'ro' | 'ru' (optional)
      ext    - 'pdf' | 'docx' | 'xlsx' (optional)
      limit  - default 10
    Returns: { items: [ {label, value(url), filename, doc_type, ext, language, code} ] }
    """
    q_raw  = (request.args.get("q") or "").strip()
    code   = _normalize_measure(request.args.get("code"))
    lang   = (request.args.get("lang") or "").strip().lower() or None
    ext    = (request.args.get("ext") or "").strip().lower() or None
    try:
        limit  = max(1, min(int(request.args.get("limit", 10)), 50))
    except Exception:
        limit = 10

    where = ["1=1"]
    params = {}

    if code:
        where.append("a.measure_code = :code")
        params["code"] = code
    if lang:
        where.append("a.language = :lang")
        params["lang"] = lang
    if ext:
        where.append("a.ext = :ext")
        params["ext"] = ext

    # Build search filter + ranking only if q is provided
    ranking_expr = None
    if q_raw:
        where.append("""(
            a.search @@ websearch_to_tsquery('simple', unaccent(:q))
            OR similarity(unaccent(a.filename), unaccent(:q)) > 0.30
            OR unaccent(a.filename) ILIKE unaccent(:like)
            OR unaccent(coalesce(a.doc_type,'')) ILIKE unaccent(:like)
        )""")
        params["q"] = q_raw
        params["like"] = f"{q_raw}%"
        ranking_expr = """greatest(
            ts_rank(a.search, websearch_to_tsquery('simple', unaccent(:q))),
            similarity(unaccent(a.filename), unaccent(:q))
        )"""

    # ORDER BY parts
    order_parts = []
    if code:
        order_parts.append("case when a.measure_code = :code then 0 else 1 end")
    order_parts.append("""
        case
          when lower(coalesce(a.doc_type,'')) like 'cerere%' then 0
          when lower(coalesce(a.doc_type,'')) like 'fișă%' or lower(coalesce(a.doc_type,'')) like 'fisa%' then 1
          when lower(coalesce(a.doc_type,'')) like 'act%' then 2
          when lower(coalesce(a.doc_type,'')) like 'anex%' then 3
          when lower(coalesce(a.doc_type,'')) like 'ordin%' then 4
          when lower(coalesce(a.doc_type,'')) like 'hotărâre%' or lower(coalesce(a.doc_type,'')) like 'hotarare%' then 5
          when lower(coalesce(a.doc_type,'')) like 'ghid%' then 6
          else 7
        end
    """)
    if ranking_expr:
        order_parts.append(f"{ranking_expr} desc")
    order_parts.append("length(coalesce(a.filename,'')) asc")

    order_clause = ",\n        ".join(order_parts)
    ranking_select = ranking_expr if ranking_expr else "0"

    sql = f"""
      select
        a.url,
        a.filename,
        a.doc_type,
        a.ext,
        a.language,
        a.measure_code as code,
        {ranking_select} as rank_score
      from public.scraped_asset a
      where {' AND '.join(where)}
      order by 
        {order_clause}
      limit :limit
    """
    params["limit"] = limit

    try:
        # bootstrap (autocommit so failures don't poison the transaction)
        with engine.connect().execution_options(isolation_level="AUTOCOMMIT") as conn:
            try:
                _bootstrap_autocomplete(conn)
            except Exception:
                pass

        # actual SELECT
        with engine.connect() as conn:
            rows = list(conn.execute(text(sql), params))
    except Exception as e:
        return jsonify(error="autocomplete_failed", details=str(e)), 500

    items = []
    for (url, filename, doc_type, ext_, language, mcode, _rank) in rows:
        left = (doc_type or "document").strip().capitalize()
        mid  = (mcode or "").replace("_", " ").strip()
        right = (filename or "").strip()
        ext_up = (ext_ or "").upper()
        label = " — ".join([p for p in [left, mid, f"{right} ({ext_up})"] if p])
        items.append({
            "label": label,
            "value": url,
            "filename": filename,
            "doc_type": doc_type,
            "ext": ext_,
            "language": language,
            "code": mcode,
        })

    return jsonify({"items": items})


# ---------- Route: POST /api/scraper/complete-docx ----------
@scraper_bp.post("/complete-docx")
def complete_docx():
    """
    Body (JSON):
    {
      // choose ONE of:
      "url": "https://…/Cerere%20SP_2.10.docx",
      "code": "SP_2.10", "doc_type": "cerere",

      // optional guidance
      "profile": { "nume_prenume": "Ion Popescu", "idno": "1007600...", ... },
      "instructions": "Ton formal, română. Completează concis.",
      "language": "ro",
      "filename": "Cerere_SP_2.10_autocomplete.docx"
    }
    Returns: a .docx download with suggestions inserted as «…»
    """
    data = request.get_json(silent=True) or {}
    filename = (data.get("filename") or "document_autocomplete.docx").strip()
    language = (data.get("language") or "ro").strip().lower()
    instructions = data.get("instructions")
    profile = data.get("profile") or {}

    try:
        with engine.begin() as conn:
            url = _resolve_source_to_url(data, conn)
        if not url:
            return jsonify(error="no_source", details="Provide 'url' or ('code' + 'doc_type') that resolves to a DOCX in scraped_asset."), 400

        if not url.lower().endswith(".docx"):
            return jsonify(error="unsupported_format", details="Only .docx is supported."), 400

        # 1) fetch source
        content = _download_bytes(url)
        doc = Document(io.BytesIO(content))

        # 2) extract fields/blanks
        fields = _extract_doc_fields(doc)

        # 3) ask AI (or fallback) for values
        suggestions = _ai_fill_fields(fields, profile=profile, instructions=instructions, language=language)

        # 4) apply inline
        _apply_suggestions_inline(doc, suggestions)

        # 5) add a small header explaining markers (first page, top)
        hdr = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        hdr_text = "«Aceste marcaje indică sugestiile auto-completate. Editați la nevoie înainte de depunere.»"
        try:
            if doc.paragraphs:
                doc.paragraphs[0].insert_paragraph_before(hdr_text)
            else:
                doc.add_paragraph(hdr_text)
        except Exception:
            # last-resort: append to top of document
            try:
                doc.add_paragraph(hdr_text)
            except Exception:
                pass

        # 6) stream out
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(
            bio,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )

    except requests.HTTPError as e:
        return jsonify(error="download_failed", details=str(e)), 502
    except Exception as e:
        return jsonify(error="autocomplete_docx_failed", details=str(e)), 500

@scraper_bp.post("/docgen")
def docgen():
    """
    JSON body (one of):
      { "url": "<direct DOCX url>" }
      { "code": "SP_2.10", "doc_type": "cerere" }

    Optional:
      "instructions": "freeform instructions for the AI",
      "lang": "ro" | "ru" | "en" (default: "ro"),
      "filename": "Optional base name for the output",
      "mode": "edit"   # (reserved; template mode can be added later)

    Returns:
      - A .docx file as an attachment (Content-Disposition: attachment)
      - On error: JSON { error, details }
    """
    data = request.get_json(force=True) or {}

    url = (data.get("url") or "").strip() or None
    code = _normalize_measure(data.get("code"))
    doc_type = _norm_doc_type(data.get("doc_type"))
    instructions = data.get("instructions") or "Curăță, structurează și optimizează conținutul pentru a fi ușor de completat."
    lang = (data.get("lang") or "ro").strip().lower()
    out_base = (data.get("filename") or "").strip() or None

    # Resolve URL from DB if needed
    if not url:
        if not (code and doc_type):
            return jsonify(error="no_source",
                           details="Provide 'url' or ('code' + 'doc_type') that resolves to a DOCX in scraped_asset."), 400
        try:
            with engine.begin() as conn:
                row = conn.execute(text("""
                    select url, filename
                      from public.scraped_asset
                     where measure_code = :code
                       and lower(coalesce(doc_type,'')) like :dt || '%%'
                       and lower(coalesce(ext,'')) = 'docx'
                     order by length(coalesce(filename,'')) asc
                     limit 1
                """), {"code": code, "dt": doc_type}).first()
        except Exception as e:
            return jsonify(error="lookup_failed", details=str(e)), 500

        if not row:
            return jsonify(error="no_match",
                           details=f"No DOCX found for {code} + {doc_type} in scraped_asset."), 404

        url, fn = row
        if not out_base:
            out_base = (fn or pathlib.Path(url).name).rsplit(".", 1)[0]

    # If caller gave URL directly
    if url and not out_base:
        out_base = pathlib.Path(url).name.rsplit(".", 1)[0]

    # Guard: DOCX only
    if not url.lower().endswith(".docx"):
        return jsonify(error="not_docx", details="Resolved URL is not a .docx"), 400

    # Pipeline
    try:
        src_bytes = _fetch_docx_bytes(url)
        src_text  = _docx_to_text(src_bytes)
        ai_text   = _call_ai_transform(src_text, instructions, lang)
        out_bytes = _write_docx_from_text(ai_text)

        out_name = re.sub(r"[^\w\-. ]+", "_", (out_base or "document")) + "_AI.docx"
        return send_file(
            io.BytesIO(out_bytes),
            as_attachment=True,
            download_name=out_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return jsonify(error="docgen_failed", details=str(e)), 500


# mount /api/scraper/*
bp.register_blueprint(scraper_bp)
