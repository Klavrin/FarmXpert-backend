import os
import re
import io
import base64
import tempfile
import pathlib
import urllib.parse
from typing import Dict, Any, Optional, List
from docx import Document
from openpyxl import load_workbook

UNDERSCORES = re.compile(r"_{4,}")  # sequences of 4+ underscores
_BLANK_RE = re.compile(r"(?:_{3,}|\.{3,}|–{3,}|—{3,})")
NBSP = "\u00A0"


def _copy_run_formatting(src_run, dst_run):
    """
    Copy a few basic formatting properties from src_run to dst_run.
    This is intentionally conservative (only common attrs) and tolerant to errors.
    """
    if src_run is None or dst_run is None:
        return

    try:
        # simple boolean properties
        for attr in ("bold", "italic", "underline", "strike"):
            if hasattr(src_run, attr) and hasattr(dst_run, attr):
                setattr(dst_run, attr, getattr(src_run, attr))
    except Exception:
        # don't crash on odd run objects
        pass

    try:
        # copy font-level attributes if available
        if hasattr(src_run, "font") and hasattr(dst_run, "font"):
            sfont = src_run.font
            dfont = dst_run.font
            try:
                if getattr(sfont, "name", None):
                    dfont.name = sfont.name
            except Exception:
                pass
            try:
                if getattr(sfont, "size", None):
                    dfont.size = sfont.size
            except Exception:
                pass
            # color may not exist or be complex; copy if present
            try:
                if getattr(sfont, "color", None) and getattr(sfont.color, "rgb", None):
                    dfont.color.rgb = sfont.color.rgb
            except Exception:
                pass
    except Exception:
        pass


def ensure_dir(p: pathlib.Path):
    p.mkdir(parents=True, exist_ok=True)

def _clean_ws(s: Optional[str]) -> str:
    return (s or "").replace(NBSP, " ").strip()

def _is_blankish(s: Optional[str]) -> bool:
    t = _clean_ws(s)
    if not t:
        return True
    if _BLANK_RE.match(t):
        return True
    return t in {".", "..", "…", "-"}

def _slug_label(s: str) -> str:
    s = _clean_ws(s).lower()
    s = re.sub(r"[^a-z0-9ăâîșț\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _pick_suggestion(label: str, suggestions: Dict[str, str]) -> Optional[str]:
    if not label:
        return None
    key = _slug_label(label)
    # Direct match
    if key in suggestions:
        return suggestions[key]
    # Substring match
    for k, v in suggestions.items():
        if k in key or key in k:
            return v
    # Fuzzy match
    words_label = set(k for k in key.split() if k)
    best = None
    best_score = 0.0
    for k, v in suggestions.items():
        toks = set(k.split())
        if not toks or not words_label:
            continue
        score = len(toks & words_label) / len(toks | words_label)
        if score > best_score:
            best_score = score
            best = v
    return best if best_score >= 0.25 else None

def _normalize_suggestions(sug: Dict[str, Any]) -> Dict[str, str]:
    out = {}
    for k, v in (sug or {}).items():
        nk = _slug_label(str(k))
        out[nk] = str(v) if v is not None else ""
    return out

def safe_filename(name: str) -> str:
    name = urllib.parse.unquote(str(name or "")).strip()
    name = re.sub(r'[\\/:*?"<>|]', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return (name[:200] or "document")


def download_url(url: str, dest: pathlib.Path) -> pathlib.Path:
    """Download a URL to a path. Gracefully handles missing 'requests'."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    try:
        import requests  # lazy import
    except Exception as e:
        raise RuntimeError(f"'requests' is not installed: {e}")

    with requests.get(url, stream=True, timeout=60) as r:
        r.raise_for_status()
        with open(dest, "wb") as f:
            for chunk in r.iter_content(chunk_size=1 << 20):
                if chunk:
                    f.write(chunk)
    return dest

def prefill_docx(src: pathlib.Path, dest: pathlib.Path, suggestions: Dict[str, Any], instructions: str | None = None, language: str = "ro"):
    """Fill blanks and label-right cells in a .docx using suggestions dict."""
    suggestions = _normalize_suggestions(suggestions)
    doc = Document(str(src))

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            for i in range(len(cells) - 1):
                left_txt = _clean_ws(cells[i].text)
                right_txt = _clean_ws(cells[i + 1].text)
                if left_txt and _is_blankish(right_txt):
                    pick = _pick_suggestion(left_txt, suggestions)
                    if pick:
                        try:
                            cells[i + 1].text = str(pick)
                        except Exception:
                            try:
                                cells[i + 1].paragraphs[0].add_run(str(pick))
                            except Exception:
                                pass

    for p in doc.paragraphs:
        text = p.text or ""
        if not text:
            continue
        m = _BLANK_RE.search(text)
        if m:
            left = text[:m.start()]
            lbl = " ".join(re.findall(r"\w+", left)[-6:])
            val = _pick_suggestion(lbl, suggestions) or _pick_suggestion(left, suggestions)
            if val:
                new_text = text[:m.end()] + f" «{val}»" + text[m.end():]
                for run in list(p.runs):
                    try:
                        run.text = ""
                    except Exception:
                        pass
                p.add_run(new_text)

    ensure_dir(dest.parent)
    doc.save(str(dest))


def prefill_xlsx(src: pathlib.Path, dest: pathlib.Path, suggestions: Dict[str, Any], instructions: str | None = None, language: str = "ro"):
    """Fill Excel cells based on left-cell labels."""
    suggestions = _normalize_suggestions(suggestions)
    wb = load_workbook(filename=str(src))
    
    for ws in wb.worksheets:
        max_row = ws.max_row
        max_col = ws.max_column
        for r in range(1, max_row + 1):
            for c in range(2, max_col + 1):
                cell = ws.cell(row=r, column=c)
                left = ws.cell(row=r, column=c - 1)
                if (cell.value is None or (isinstance(cell.value, str) and not cell.value.strip())) and left and left.value:
                    label = str(left.value)
                    pick = _pick_suggestion(label, suggestions)
                    if pick:
                        try:
                            cell.value = pick
                        except Exception:
                            pass

    ensure_dir(dest.parent)
    wb.save(str(dest))


# ----------------------------- OPTIONAL HELPERS -----------------------------

def _first(d: dict, keys: List[str], default=None):
    """Safely fetch the first present & non-empty key."""
    for k in keys:
        if isinstance(d, dict) and k in d and d[k] not in (None, "", []):
            return d[k]
    return default


def _join_nonempty(*vals, sep=" "):
    return sep.join(str(v) for v in vals if v not in (None, "", []))


def _build_suggestions(profile: dict, subsidy: Optional[dict] = None, lang: str = "ro") -> Dict[str, str]:
    """
    Produce common AIPA fields. Tolerant to missing pieces in 'profile'.
    'profile' can be whatever your farm_profile/service returns.
    """
    user = profile.get("user") or profile.get("utilizator") or {}
    business = profile.get("business") or {}
    finance = profile.get("finance") or profile.get("finante") or {}
    fields = profile.get("fields") or profile.get("campuri") or []
    animals = profile.get("animals") or profile.get("efective_animale") or []
    vehicles = profile.get("vehicles") or profile.get("machines") or []
    cattle = profile.get("cattle") or []

    first_name = _first(user, ["firstName", "first_name", "firstname", "nume"])
    last_name = _first(user, ["lastName", "last_name", "lastname", "prenume"])
    solicitant = _join_nonempty(first_name, last_name) or "[Nume solicitant]"
    idno = _first(business, ["idno", "id", "businessId", "business_id"], _first(user, ["businessId", "business_id"]))
    phone = _first(user, ["phone", "telefon"])
    email = _first(user, ["email"])
    denumire = _first(business, ["name", "denumire", "title"], solicitant)

    # fields aggregation
    total_ha = 0.0
    crop_types: List[str] = []
    for f in fields:
        try:
            if f.get("size") is not None:
                total_ha += float(f["size"])
        except Exception:
            pass
        ct = f.get("cropType") or f.get("crop_type")
        if ct:
            crop_types.append(str(ct))
    crop_types_txt = ", ".join(sorted(set(crop_types))) if crop_types else "—"

    # animals aggregation
    animal_counts: Dict[str, int] = {}
    for a in animals:
        sp = (a.get("species") or a.get("type") or "").strip().lower() or "necunoscut"
        animal_counts[sp] = animal_counts.get(sp, 0) + int(a.get("amount") or 1)
    animale_txt = ", ".join(f"{k}: {v}" for k, v in animal_counts.items()) if animal_counts else "—"

    # vehicles
    nr_veh = len(vehicles) if isinstance(vehicles, list) else 0

    # bovines if provided
    bovine = 0
    try:
        for c in cattle:
            amt = c.get("amount")
            if amt is not None:
                bovine += int(amt)
    except Exception:
        pass

    venit_anual = _first(finance, ["yearlyIncome", "inc"])
    chelt_anual = _first(finance, ["yearlyExpenses", "exp"])

    su_code = (subsidy or {}).get("code") or ""
    su_title = (subsidy or {}).get("title") or ""

    suggestions = {
        "Denumirea solicitantului": denumire or solicitant,
        "Numele și prenumele": solicitant,
        "IDNO/Cod fiscal": idno or "[IDNO]",
        "Telefon": phone or "[Telefon]",
        "Email": email or "[Email]",
        "Măsura/Program": _join_nonempty(su_code, su_title, sep=" – ") if (su_code or su_title) else "—",
        "Suprafața totală exploatată (ha)": f"{total_ha:.2f}",
        "Culturile principale": crop_types_txt,
        "Efective de animale": animale_txt,
        "Număr vehicule agricole": str(nr_veh),
        "Efectiv bovine (capete)": str(bovine),
        "Localitate": business.get("localitate") or business.get("address") or "[Localitate]",
        "Data completării": "[Auto-completare la tipărire]",
    }
    if venit_anual is not None:
        suggestions["Venit anual (lei)"] = str(venit_anual)
    if chelt_anual is not None:
        suggestions["Cheltuieli anuale (lei)"] = str(chelt_anual)

    return suggestions


def _build_plan_5_ani(profile: dict, subsidy: Optional[dict] = None, lang: str = "ro") -> str:
    user = profile.get("user") or profile.get("utilizator") or {}
    business = profile.get("business") or {}
    name = _join_nonempty(_first(user, ["firstName", "nume"]), _first(user, ["lastName", "prenume"])) or (business.get("name") or "Solicitant")
    su = (subsidy or {})
    return (
        f"Plan investițional pe 5 ani – {name}\n"
        f"Măsura: {su.get('code','—')} – {su.get('title','—')}\n\n"
        "Anul 1: Achiziție/implementare echipamente și inițiere lucrări. "
        "Întărirea capacității de producție, instruire personal, setarea evidenței costurilor.\n"
        "Anul 2: Optimizarea fluxurilor, creșterea productivității, extinderea suprafețelor/echipamentelor după necesar.\n"
        "Anul 3: Diversificare produse/servicii, îmbunătățire calitate, certificări dacă e cazul.\n"
        "Anul 4: Scalare și integrare verticală (prelucrare/comercializare), parteneriate noi.\n"
        "Anul 5: Consolidare poziție pe piață, evaluare rezultate, plan pentru reinvestirea profitului.\n\n"
        "Indicatori: productivitate/ha, randament utilaje, venit operațional, reducerea cheltuielilor, calitate.\n"
        "Riscuri: climă (irigare/tehnologii adecvate), prețuri (contracte forward), operațional (mentenanță, asigurări).\n"
    )


def prepare_documents(profile: dict, docs: List[dict], subsidy: Optional[dict] = None, lang: str = "ro") -> dict:
    """
    docs: [{ "url": "https://.../Cerere_SP_2.10.docx" }, ...]
    Returns:
       {
         "files": [{"sourceUrl","filename","ext","b64"}],
         "generated": {"plan_5_ani": "..."}
       }
    """
    suggestions = _build_suggestions(profile, subsidy, lang=lang)
    out_files: List[dict] = []

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = pathlib.Path(tmpdir)
        for item in (docs or []):
            url = item.get("url")
            if not url:
                continue
            src_name = safe_filename(url.split("/")[-1])
            ext = pathlib.Path(src_name).suffix.lower().lstrip(".")
            src = tmp / src_name
            download_url(url, src)

            dest_name = f"{pathlib.Path(src_name).stem}_precompletat.{ext or 'docx'}"
            dest = tmp / dest_name

            if ext == "docx":
                prefill_docx(src, dest, suggestions)
            elif ext == "xlsx":
                prefill_xlsx(src, dest, suggestions)
            else:
                # Unknown type: pass-through original
                dest = src

            b = dest.read_bytes()
            b64 = base64.b64encode(b).decode("ascii")
            out_files.append({
                "sourceUrl": url,
                "filename": dest_name,
                "ext": ext,
                "b64": b64,
            })

    plan = _build_plan_5_ani(profile, subsidy, lang=lang)
    return {"files": out_files, "generated": {"plan_5_ani": plan}}
